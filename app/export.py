import re
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def save_as_docx(markdown_text: str, output_path: Path, title: str = ""):
    """Convert markdown notes to a formatted .docx file."""
    doc = Document()

    if title:
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for line in markdown_text.splitlines():
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif re.match(r"^    [-*] |^  {2,4}[-*] ", line):
            # Indented bullet → sub-bullet
            text = re.sub(r"^[ \t]+[-*] ", "", line)
            p = doc.add_paragraph(style="List Bullet 2")
            _add_inline_formatting(p, text.strip())
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_formatting(p, line[2:].strip())
        elif line.startswith("> "):
            p = doc.add_paragraph(style="Quote")
            _add_inline_formatting(p, line[2:].strip())
        elif line.strip() == "" or line.strip() == "---":
            doc.add_paragraph()
        else:
            p = doc.add_paragraph()
            _add_inline_formatting(p, line.strip())

    doc.save(str(output_path))


def _add_inline_formatting(paragraph, text: str):
    """Parse **bold** and *italic* and render them in paragraph runs."""
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)
